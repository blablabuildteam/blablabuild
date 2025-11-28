#!/usr/bin/env python3
"""
Script to extract all text from the homepage and generate a Word document.
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_homepage_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('blablabuild Homepage - Complete Text Content', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Navigation Section
    doc.add_heading('Navigation', level=1)
    doc.add_paragraph('blablabuild')
    doc.add_paragraph('Aanpak')
    doc.add_paragraph('Team')
    doc.add_paragraph('Impact')
    doc.add_paragraph('Cases')
    doc.add_paragraph('Gratis AI Advies')
    
    # Hero Section
    doc.add_heading('Hero Section', level=1)
    doc.add_paragraph('minder praten, meer bouwen')
    
    hero_heading = doc.add_paragraph()
    hero_heading.add_run('Een chat. ').bold = True
    hero_heading.add_run('Een meeting. ').bold = True
    hero_heading.add_run('Directe AI Impact').bold = True
    
    doc.add_paragraph('Bespaar kostbare tijd met slimme automatisering')
    doc.add_paragraph('Reduceer operationele kosten met AI-gedreven efficiency')
    doc.add_paragraph('Verhoog je omzet door data-driven beslissingen')
    doc.add_paragraph('Elimineer wrijving in je processen')
    doc.add_paragraph('Centraliseer je data voor betere inzichten')
    
    # Approach Section
    doc.add_heading('Aanpak (Approach) Section', level=1)
    doc.add_paragraph('VAN EERSTE CONTACT TOT SCHAALBARE IMPACT')
    
    approach_heading = doc.add_paragraph()
    approach_heading.add_run('Geen agency bullsh').bold = True
    approach_heading.add_run('*').bold = True
    approach_heading.add_run('t, simpelweg resultaat').bold = True
    
    doc.add_paragraph('Een gestructureerde aanpak die van eerste contact tot schaalbare impact leidt.')
    
    doc.add_heading('Stap 1: Bla', level=2)
    doc.add_paragraph('We bellen een keer of doen een koffietje om jouw situatie te bespreken')
    
    doc.add_heading('Stap 2: Bla', level=2)
    doc.add_paragraph('We bereiden een sessie voor om met jou en je team de diepte in te duiken en een plan te maken.')
    
    doc.add_heading('Stap 3: Build', level=2)
    doc.add_paragraph('We gaan direct aan de slag om in enkele weken impact te leveren.')
    
    # Team Section
    doc.add_heading('Team Section', level=1)
    doc.add_paragraph('Senioriteit zonder Overhead')
    
    doc.add_heading('Het High-Impact Orchestration Team', level=2)
    
    doc.add_heading('Daniel - Data, Tech & AI', level=3)
    doc.add_paragraph('Focus: AI, Technologie en Data')
    doc.add_paragraph('Brengt strategie, data en cutting-edge AI-technologie samen. Vertaalt complexe uitdagingen naar slimme, schaalbare oplossingen door razendsnelle prototyping.')
    doc.add_paragraph('• Toekomstbestendige AI-Strategie')
    doc.add_paragraph('• Operationele AI/Data Workflows')
    doc.add_paragraph('• Bewezen Thought Leadership')
    doc.add_paragraph('• Prototyping Expert')
    
    doc.add_heading('Kevin - Growth & CX', level=3)
    doc.add_paragraph('Focus: Markt, Merk en Conversie')
    doc.add_paragraph('Combineert strategische visie met hands-on ondernemerschap om schaalbare digitale oplossingen te leveren. Specialisatie ligt in het winnen van de markt door een sterke merkidentiteit en conversiekracht.')
    doc.add_paragraph('• E-commerce & Conversie')
    doc.add_paragraph('• Merkopbouw & Emotie')
    doc.add_paragraph('• Data-gedreven Groei')
    
    doc.add_heading('Xennith - Business Transformation', level=3)
    doc.add_paragraph('Focus: Structuur, Proces & Implementatie')
    doc.add_paragraph('Combineert AI consulting, tech en productie kennis om complexiteit te vertalen naar concrete en uitvoerbare kansen met focus op het stroomlijnen organisaties.')
    doc.add_paragraph('• Enterprise Strategie & Ervaring')
    doc.add_paragraph('• Van Pijn naar Plan')
    doc.add_paragraph('• Meetbaar Groei Focus')
    doc.add_paragraph('• Operationele Efficiëntie')
    
    doc.add_paragraph()
    doc.add_paragraph('Gecombineerd meer dan 50 jaar digitale ervaring ― nu beschikbaar voor jouw innovaties')
    
    # Impact Section
    doc.add_heading('Impact Section', level=1)
    doc.add_paragraph('GEGARANDEERDE RESULTATEN')
    
    doc.add_heading('Onze Impact', level=2)
    
    doc.add_heading('01 - Data & AI-Strategie', level=3)
    doc.add_paragraph('We vertalen complexe data naar een duidelijk overzicht van jouw kansen.')
    
    doc.add_heading('02 - High-Impact Groei', level=3)
    doc.add_paragraph('We verbeteren de klantervaring voor snelle impact en extra omzet.')
    
    doc.add_heading('03 - Automatisering & Efficiëntie', level=3)
    doc.add_paragraph('We zetten onze senioriteit en AI-engine in om processen te versnellen en tijd te besparen.')
    
    doc.add_paragraph('Gratis AI Advies')
    
    # Cases Section
    doc.add_heading('Cases Section', level=1)
    doc.add_paragraph('BEWEZEN RESULTATEN')
    
    doc.add_heading('Cases', level=2)
    doc.add_paragraph('Ontdek hoe we impact hebben geleverd voor onze klanten')
    
    doc.add_heading('Case 1: E-commerce Platform Transformation', level=3)
    doc.add_paragraph('Badge: Data & AI')
    doc.add_paragraph('• 80% reductie in operationele kosten')
    doc.add_paragraph('• 3x snellere order processing')
    doc.add_paragraph('• Real-time inventory tracking')
    doc.add_paragraph('We hebben een volledig geautomatiseerd e-commerce platform gebouwd dat real-time inventory tracking combineert met AI-gedreven voorspellingen. Het resultaat: een 80% reductie in operationele kosten en 3x snellere order processing. De oplossing integreert naadloos met bestaande systemen en schaalt automatisch mee met de groei van het bedrijf.')
    doc.add_paragraph('Klik om meer te lezen')
    doc.add_paragraph('Klik om terug te gaan')
    
    doc.add_heading('Case 2: Lead Generation Automatisering', level=3)
    doc.add_paragraph('Badge: Growth & CX')
    doc.add_paragraph('• 250% verhoging in kwalitatieve leads')
    doc.add_paragraph('• 60% tijd bespaard op lead qualification')
    doc.add_paragraph('• AI-powered lead scoring systeem')
    doc.add_paragraph('Door het implementeren van een geavanceerd AI-systeem voor lead qualification en scoring, hebben we de kwaliteit van leads met 250% verhoogd. Het systeem bespaart het sales team 60% van hun tijd door automatisch leads te scoren en te categoriseren op basis van gedrag en intentie.')
    doc.add_paragraph('Klik om meer te lezen')
    doc.add_paragraph('Klik om terug te gaan')
    
    doc.add_heading('Case 3: Supply Chain Optimalisatie', level=3)
    doc.add_paragraph('Badge: Automatisering')
    doc.add_paragraph('• 40% reductie in voorraadkosten')
    doc.add_paragraph('• Real-time tracking & voorspellingen')
    doc.add_paragraph('• Geautomatiseerde bestelprocessen')
    doc.add_paragraph('Een volledig geautomatiseerd supply chain management systeem dat real-time tracking combineert met voorspellende analytics. Het systeem heeft geleid tot een 40% reductie in voorraadkosten door slimmere bestelprocessen en nauwkeurige vraagvoorspellingen. Alle processen zijn geautomatiseerd, van bestelling tot levering.')
    doc.add_paragraph('Klik om meer te lezen')
    doc.add_paragraph('Klik om terug te gaan')
    
    doc.add_heading('Case 4: Customer Analytics Dashboard', level=3)
    doc.add_paragraph('Badge: Data & AI')
    doc.add_paragraph('• 360° klantbeeld in real-time')
    doc.add_paragraph('• Voorspellende customer insights')
    doc.add_paragraph('• Geautomatiseerde rapportage')
    doc.add_paragraph('Een geavanceerd analytics dashboard dat alle klantdata centraliseert en transformeert naar actionable insights. Het systeem biedt een 360° klantbeeld in real-time en gebruikt AI om voorspellende analyses te maken. Automatische rapportage bespaart weken aan handmatig werk.')
    doc.add_paragraph('Klik om meer te lezen')
    doc.add_paragraph('Klik om terug te gaan')
    
    doc.add_heading('Case 5: Personalized Marketing Platform', level=3)
    doc.add_paragraph('Badge: Growth & CX')
    doc.add_paragraph('• 35% verhoging in conversie')
    doc.add_paragraph('• Geautomatiseerde personalisatie')
    doc.add_paragraph('• Cross-channel campagne management')
    doc.add_paragraph('Een geïntegreerd marketing platform dat AI gebruikt om elke klantinteractie te personaliseren. Het systeem heeft geleid tot een 35% verhoging in conversie door relevante content en aanbiedingen op het juiste moment te leveren. Alle campagnes worden automatisch geoptimaliseerd en beheerd.')
    doc.add_paragraph('Klik om meer te lezen')
    doc.add_paragraph('Klik om terug te gaan')
    
    doc.add_heading('Case 6: Workflow Automatisering Suite', level=3)
    doc.add_paragraph('Badge: Automatisering')
    doc.add_paragraph('• 70% reductie in handmatig werk')
    doc.add_paragraph('• Geautomatiseerde document processing')
    doc.add_paragraph('• Seamless integratie met bestaande tools')
    doc.add_paragraph('Een uitgebreide suite van geautomatiseerde workflows die repetitieve taken elimineert. Het systeem heeft 70% van het handmatige werk geautomatiseerd, inclusief document processing, data entry en communicatie. Alle workflows integreren naadloos met bestaande tools en systemen.')
    doc.add_paragraph('Klik om meer te lezen')
    doc.add_paragraph('Klik om terug te gaan')
    
    # CTA Section
    doc.add_heading('CTA Section', level=1)
    doc.add_heading('Klaar om te starten?', level=2)
    doc.add_paragraph('Beantwoord 7 vragen en ontvang binnen 5 minuten een gepersonaliseerde AI-analyse.')
    doc.add_paragraph('Start je gratis analyse')
    
    # Footer
    doc.add_heading('Footer', level=1)
    doc.add_paragraph('build')
    doc.add_paragraph('© 2025 blablabuild')
    doc.add_paragraph('hello@blablabuild.com')
    
    # Save document
    output_path = 'homepage-text-content.docx'
    doc.save(output_path)
    print(f'Word document created successfully: {output_path}')

if __name__ == '__main__':
    create_homepage_document()

